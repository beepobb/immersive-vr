from docx import Document


def save_transcript_to_docx(
    transcript_text: str,
    output_path: str,
    title: str = "Transcript"
) -> str:
    """
    Create a .docx file containing the transcript text.

    Args:
        transcript_text: The final transcript text.
        output_path: Where to save the .docx file.
        title: Heading shown at the top of the document.

    Returns:
        The output path of the saved document.
    """
    
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(transcript_text)
    document.save(output_path)
    return output_path